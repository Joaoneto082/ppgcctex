# -*- coding: utf-8 -*-
"""
Gera um arquivo Word (.docx) com o cronograma de escrita da dissertacao
de Joao Marinho Neto - Mestrado em Ciencia da Computacao (UFERSA/UERN).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- helpers ----------
def set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading_blue(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
        run.font.name = "Calibri"


def add_para(doc: Document, text: str, bold: bool = False, size: int = 11,
             italic: bool = False, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_bullets(doc: Document, items) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def make_table(doc: Document, headers, rows, widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        set_cell_bg(hdr[i], "1F3A68")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(r):
            row_cells[i].text = ""
            p = row_cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if widths_cm:
        for row in table.rows:
            for i, w in enumerate(widths_cm):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table


# ---------- documento ----------
doc = Document()

# margens
for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# estilo padrao
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ---------- capa ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Cronograma de Escrita da Dissertação")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
run.font.name = "Calibri"

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run(
    "Análise Automatizada de Objetivos de Desenvolvimento Sustentável: "
    "Uma Abordagem Baseada em Processamento de Linguagem Natural e Modelos de "
    "Linguagem de Grande Porte Aplicada a Mossoró-RN"
)
run.italic = True
run.font.size = Pt(12)
run.font.name = "Calibri"

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(
    "Mestrando: João Marinho Neto\n"
    "Orientador: Prof. Dr. Francisco Milton Mendes Neto\n"
    "PPGCC – UFERSA/UERN\n"
    "Data de elaboração do cronograma: 30/04/2026"
)
run.font.size = Pt(11)
run.font.name = "Calibri"

doc.add_paragraph()

# ---------- 1. Diagnostico ----------
add_heading_blue(doc, "1. Diagnóstico do estado atual", level=1)

add_heading_blue(doc, "1.1 O que já está pronto na dissertação (documento.pdf)", level=2)
add_bullets(doc, [
    "Pré-textuais (resumo + abstract com foco em ODS Mossoró – OK).",
    "Cap. 1 Justificativa, Cap. 2 Introdução, Cap. 3 Hipótese, Cap. 4 Objetivos.",
    "Cap. 5 Revisão Sistemática da Literatura.",
    "Cap. 6 Materiais e Métodos.",
    "Cap. 7 Implementação (3 versões, módulo de coleta, filtragem e IA).",
])

add_heading_blue(
    doc,
    "1.2 O que já existe na implementação mas NÃO está no documento.pdf",
    level=2,
)
add_bullets(doc, [
    "ODS 3 (Saúde) e ODS 4 (Educação), além de ODS 9 e 16.",
    "Replicação geográfica: Natal-RN, Fortaleza-CE, João Pessoa-PB, Salvador-BA.",
    "Pasta ODS_generica/ com ods_analise_generica.py + formulario_config_ods.html "
    "(gerador interativo de configurações).",
    "Dashboards web em public/ (dashboard_ods16.html, diagnostico_ods9.html, etc.) "
    "hospedados em dashboard-ods-mossoro.netlify.app — hoje a dissertação só cita Power BI.",
])

add_heading_blue(
    doc,
    "1.3 O que já existe FORA da dissertação e ainda não foi incorporado",
    level=2,
)
add_bullets(doc, [
    "Artigo do EATIS 2026 (Analise_ODS_Mossoro__Versao_Final_.pdf) com fundamentação "
    "teórica nova, trabalhos relacionados, resultados quantitativos (253 análises ODS 9 / "
    "116 análises ODS 16) e referências [1-24] mais maduras.",
    "Validação TAM aplicada – 4 respostas (Roseano Medeiros/DGA-UERN, Diego Tobias/PREVI "
    "Mossoró, Gabriella Vitorino/UFERSA-Angicos, Carlos Silva/DCAF-UFERSA) com PU, PFU, "
    "ICU e perguntas abertas.",
    "Tudo isso precisa virar capítulo/seção na dissertação.",
])

add_heading_blue(doc, "1.4 Inconsistências críticas a corrigir", level=2)
add_bullets(doc, [
    "Capítulos 1–4 (justificativa, introdução, hipótese, objetivos específicos) ainda "
    "mencionam plataforma Gnomon, identidade de lugar e sentimento de pertencimento — "
    "resíduo da qualificação em conflito direto com o resumo, abstract e Cap. 7 atuais.",
    "Em documento.tex falta \\input{} para os capítulos 8-resultados, 9-discussao e "
    "10-conclusao (os arquivos existem em 2-textuais/ mas não estão sendo compilados).",
    "Falta capítulo (ou seção) de Trabalhos Relacionados separado da RSL — o EATIS já "
    "traz esse texto pronto.",
])

doc.add_paragraph()

# ---------- 2. Cronograma ----------
add_heading_blue(doc, "2. Cronograma de escrita para defesa", level=1)

prem = doc.add_paragraph()
run = prem.add_run("Premissa: ")
run.bold = True
run.font.name = "Calibri"
run = prem.add_run(
    "hoje é 30/04/2026 e o objetivo é entregar a versão final + defender. "
    "Estão propostas ~10 semanas de escrita ativa + 2 semanas de banca, "
    "totalizando uma defesa entre meados de julho e início de agosto/2026. "
    "Ajustar as datas conforme o prazo real com o orientador."
)
run.font.name = "Calibri"

# ---- visao geral ----
add_heading_blue(doc, "2.1 Visão geral por semana", level=2)

vg_headers = ["Semana", "Período", "Foco principal"]
vg_rows = [
    ["1",  "01–07/Maio",       "Realinhamento textual dos capítulos antigos (Justif./Introd./Hipótese/Objetivos)"],
    ["2",  "08–14/Maio",       "Fundamentação Teórica + Trabalhos Relacionados"],
    ["3",  "15–21/Maio",       "Atualização da Implementação (Cap. 7) – ODS 3/4, generalização, formulário, dashboards web"],
    ["4",  "22–28/Maio",       "Resultados (parte 1) – ODS 9 e ODS 16"],
    ["5",  "29/Mai–04/Junho",  "Resultados (parte 2) – ODS 3, ODS 4 e replicação geográfica"],
    ["6",  "05–11/Junho",      "Validação TAM (capítulo dedicado)"],
    ["7",  "12–18/Junho",      "Discussão"],
    ["8",  "19–25/Junho",      "Conclusão e ajustes finos (referências, listas, apêndices)"],
    ["9",  "26/Junho–02/Julho","Revisão integral e formatação ABNT"],
    ["10", "03–09/Julho",      "Versão para a banca + entrega ao orientador"],
    ["11–12", "10–25/Julho",   "Defesa: depósito, apresentação e banca"],
]
make_table(doc, vg_headers, vg_rows, widths_cm=[1.5, 3.5, 11.0])

doc.add_paragraph()

# ---- semanas detalhadas ----
def add_week(num, periodo, titulo, headers, rows, widths):
    add_heading_blue(doc, f"Semana {num} ({periodo}) — {titulo}", level=2)
    make_table(doc, headers, rows, widths_cm=widths)
    doc.add_paragraph()


# Semana 1
add_week(
    1, "01–07/Maio",
    "Realinhamento textual dos capítulos antigos",
    ["#", "Tarefa", "Arquivo", "Esforço"],
    [
        ["1.1", "Reescrever Justificativa removendo Gnomon/identidade-de-lugar e ancorando em monitoramento de ODS municipais (usar §1 do EATIS como base)", "2-textuais/1-justificativa.tex", "1 dia"],
        ["1.2", "Reescrever Introdução (contexto ODS, dados perceptivos georreferenciados, Smart Cities)", "2-textuais/2-introducao.tex", "1 dia"],
        ["1.3", "Reescrever Hipótese para o novo escopo (LLMs + Google Places revelam gargalos de ODS municipais não capturados em indicadores oficiais)", "2-textuais/3-hipotese-objetivos.tex", "0,5 dia"],
        ["1.4", "Reescrever Objetivo Geral e Específicos (remover sentimento de pertencimento; trocar por: detectar problemas por aspecto do ODS, gerar diagnósticos, generalizar para múltiplos ODS/municípios, validar com TAM)", "2-textuais/3-hipotese-objetivos.tex", "0,5 dia"],
        ["1.5", "Atualizar documento.tex incluindo novos \\input{} (cap. 6, 8, 9, 10) e remover o cap. 6-cronograma resíduo da qualificação", "documento.tex", "0,5 dia"],
    ],
    [1.0, 8.5, 5.0, 1.7],
)

# Semana 2
add_week(
    2, "08–14/Maio",
    "Fundamentação Teórica + Trabalhos Relacionados",
    ["#", "Tarefa", "Arquivo", "Esforço"],
    [
        ["2.1", "Criar capítulo de Fundamentação Teórica autônomo (ODS, ODS 9/16/3/4, mensuração local, PLN, análise de sentimentos, Transformers, LLMs) — basear em §2 do EATIS expandindo", "2-textuais/2-fundamentacao-teorica.tex (novo)", "3 dias"],
        ["2.2", "Criar capítulo de Trabalhos Relacionados (Devlin, Gupta, Schimanski, Jaiswal, Liu, Bhattacherjee, Lee + lacuna preenchida) — adaptar §3 do EATIS", "2-textuais/6-trabalhos-relacionados.tex", "2 dias"],
        ["2.3", "Mesclar conteúdo Materiais e Métodos > Transformers/BERT com a nova Fundamentação Teórica para evitar duplicação", "2-textuais/5-materiais-metodos.tex", "0,5 dia"],
    ],
    [1.0, 8.5, 5.0, 1.7],
)

# Semana 3
add_week(
    3, "15–21/Maio",
    "Atualização da Implementação (Cap. 7)",
    ["#", "Tarefa", "Arquivo", "Esforço"],
    [
        ["3.1", "Adicionar seção 7.7 Generalização para múltiplos ODS (ODS 3, ODS 4) com configurações ods3_config.json e ods4_config.json", "2-textuais/7-implementacao.tex", "1 dia"],
        ["3.2", "Adicionar seção 7.8 Generalização geográfica (Natal, Fortaleza, João Pessoa, Salvador) — mostrar arquivos *_analise_*.json em outras cidades", "2-textuais/7-implementacao.tex", "1 dia"],
        ["3.3", "Adicionar seção 7.9 Formulário interativo de configuração (formulario_config_ods.html + ods_analise_generica.py) com screenshot e fluxo de uso", "2-textuais/7-implementacao.tex", "1 dia"],
        ["3.4", "Adicionar seção 7.10 Dashboards Web Públicos — substituir/complementar a referência a Power BI pelo dashboard hospedado em Netlify (dashboard-ods-mossoro.netlify.app)", "2-textuais/7-implementacao.tex", "1 dia"],
        ["3.5", "Atualizar diagrama de arquitetura geral (Figura 1 do Cap. 7) para incluir formulário HTML + dashboard web", "figuras/", "0,5 dia"],
    ],
    [1.0, 8.5, 5.0, 1.7],
)

# Semana 4
add_week(
    4, "22–28/Maio",
    "Resultados (parte 1) – ODS 9 e ODS 16",
    ["#", "Tarefa", "Arquivo", "Esforço"],
    [
        ["4.1", "Resultados ODS 9 (253 análises, 117 lugares, 56,9% problemas, distribuição por dimensão, tabelas/figuras) — adaptar §6 do EATIS expandindo a análise por dimensão", "2-textuais/8-resultados.tex", "2 dias"],
        ["4.2", "Resultados ODS 16 (116 análises, 12 estabelecimentos, 46,6%, dimensão Instituições Eficazes 37,1%) — adaptar §6 do EATIS", "2-textuais/8-resultados.tex", "2 dias"],
        ["4.3", "Inserir gráficos do EATIS em alta resolução (Figura 1 — distribuição de sentimentos comparada; Figura 2 — problemas por dimensão ODS 16)", "figuras/", "1 dia"],
    ],
    [1.0, 8.5, 5.0, 1.7],
)

# Semana 5
add_week(
    5, "29/Mai–04/Junho",
    "Resultados (parte 2) – ODS 3, ODS 4 e generalização",
    ["#", "Tarefa", "Arquivo", "Esforço"],
    [
        ["5.1", "Resultados ODS 3 — Mossoró e Fortaleza (a partir de analises/saude/*.json)", "2-textuais/8-resultados.tex", "1,5 dia"],
        ["5.2", "Resultados ODS 4 — Natal-RN e Fortaleza-CE (a partir de analises/educacao/*.json)", "2-textuais/8-resultados.tex", "1,5 dia"],
        ["5.3", "Resultados de replicação geográfica (Salvador, João Pessoa para ODS 16; Natal, Fortaleza para ODS 9) — comparativo entre cidades", "2-textuais/8-resultados.tex", "1 dia"],
        ["5.4", "Diagnósticos automáticos gerados pela IA — exemplos representativos (1 por ODS)", "2-textuais/8-resultados.tex", "1 dia"],
    ],
    [1.0, 8.5, 5.0, 1.7],
)

# Semana 6
add_week(
    6, "05–11/Junho",
    "Validação TAM (capítulo dedicado)",
    ["#", "Tarefa", "Arquivo", "Esforço"],
    [
        ["6.1", "Subseção em Materiais e Métodos: descrever o instrumento TAM (Davis, 1989) — PU, PFU, ICU, escala Likert, perguntas abertas, link do dashboard avaliado", "2-textuais/5-materiais-metodos.tex", "1 dia"],
        ["6.2", "Capítulo/seção de Validação TAM: perfil dos 4 respondentes (Roseano/DGA-UERN, Diego/Previ-Mossoró, Gabriella/UFERSA-Angicos, Carlos/DCAF-UFERSA)", "2-textuais/8-resultados.tex (seção) ou novo 2-textuais/9-validacao-tam.tex", "1 dia"],
        ["6.3", "Tabular respostas (PU1, PU2, PU3, PFU1, PFU2, PFU3, ICU1, ICU2 + escalas 1-5 + ranqueamento) e calcular médias por construto", "mesmo arquivo", "1 dia"],
        ["6.4", "Análise qualitativa das respostas abertas (sugestão de filtros tipo Excel, excelente iniciativa, nenhum problema encontrado) + ameaças à validade do TAM com n=4", "mesmo arquivo", "1 dia"],
        ["6.5", "(Recomendado) Coletar mais 4–6 respostas para fortalecer o TAM antes da defesa — TAM com n=4 é frágil para a banca", "em paralelo", "semanas 6–8"],
    ],
    [1.0, 8.0, 5.5, 1.7],
)

# Semana 7
add_week(
    7, "12–18/Junho",
    "Discussão",
    ["#", "Tarefa", "Arquivo", "Esforço"],
    [
        ["7.1", "Discussão dos resultados quantitativos: por que ODS 9 tem mais neutros (normalização de problemas crônicos) — usar §7 do EATIS", "2-textuais/9-discussao.tex", "1,5 dia"],
        ["7.2", "Discussão da generalização: o que mudou em outras cidades; estabilidade dos padrões", "2-textuais/9-discussao.tex", "1 dia"],
        ["7.3", "Discussão da aceitação tecnológica (TAM): triangulação com o resultado quantitativo dos dashboards", "2-textuais/9-discussao.tex", "1 dia"],
        ["7.4", "Limitações: viés de seleção/digital, cobertura desigual, temporalidade, viés de termos de busca, n pequeno do TAM (§7 do EATIS já traz)", "2-textuais/9-discussao.tex", "1 dia"],
        ["7.5", "Implicações para Smart Cities e gestão pública municipal", "2-textuais/9-discussao.tex", "0,5 dia"],
    ],
    [1.0, 8.5, 5.0, 1.7],
)

# Semana 8
add_week(
    8, "19–25/Junho",
    "Conclusão e ajustes finos",
    ["#", "Tarefa", "Arquivo", "Esforço"],
    [
        ["8.1", "Conclusão: retomada da hipótese, contribuições, resposta à QGP e às QC/QT/QP da Design Science", "2-textuais/10-conclusao.tex", "1,5 dia"],
        ["8.2", "Trabalhos futuros (4 eixos do EATIS: temático, geográfico, longitudinal, melhorias metodológicas, integração com outros dados, plataforma web para gestores)", "2-textuais/10-conclusao.tex", "0,5 dia"],
        ["8.3", "Atualizar referencias.bib com toda a bibliografia do EATIS (Abidoye, Brodny, Devlin, Fu, Hu, Jaiswal, Larosa, Liu, Massey, Parvathy, Schimanski, Wankhade, Água, etc.)", "3-pos-textuais/referencias.bib", "1 dia"],
        ["8.4", "Atualizar lista de figuras, lista de tabelas, lista de abreviaturas (incluir TAM, PU, PFU, ICU, IDSC)", "pré-textuais", "0,5 dia"],
        ["8.5", "Apêndices: questionário TAM (form), arquivos de configuração JSON exemplares, link do repositório/Zenodo", "3-pos-textuais/apendices/", "1 dia"],
    ],
    [1.0, 8.5, 5.0, 1.7],
)

# Semana 9
add_week(
    9, "26/Junho–02/Julho",
    "Revisão integral e formatação ABNT",
    ["#", "Tarefa", "Esforço"],
    [
        ["9.1", "Leitura crítica do texto inteiro de ponta a ponta (consistência terminológica: ODS/SDG, Mossoró-RN/Mossoró, LLM/modelos de linguagem)", "2 dias"],
        ["9.2", "Revisão ortográfica e gramatical", "1 dia"],
        ["9.3", "Verificar normas ABNT (citações, referências, sumário, listas)", "1 dia"],
        ["9.4", "Compilar PDF final, gerar índices/glossário, conferir paginação", "0,5 dia"],
    ],
    [1.0, 12.5, 2.5],
)

# Semana 10
add_heading_blue(
    doc, "Semana 10 (03–09/Julho) — Versão para a banca + entrega ao orientador", level=2
)
add_bullets(doc, [
    "Entregar ao orientador na segunda-feira 06/07 com pelo menos 1 semana para "
    "ele revisar antes do depósito formal.",
    "Ajustes finais conforme feedback do Prof. Milton.",
])
doc.add_paragraph()

# Semanas 11-12
add_heading_blue(doc, "Semanas 11–12 (10–25/Julho) — Defesa", level=2)
add_bullets(doc, [
    "Depósito da dissertação com a banca (Prof. Milton + Prof. Alex Sandro Gomes/UFPE "
    "+ Prof. Bruno de Sousa Monteiro/UFERSA).",
    "Preparar apresentação (~25 slides, 30 min): contexto → metodologia → implementação "
    "(com demo do dashboard ao vivo) → resultados (ODS 9, 16, 3, 4 e replicação) → TAM "
    "→ contribuições.",
    "Defesa estimada: ~25/Julho – 01/Agosto/2026.",
])
doc.add_paragraph()

# ---------- 3. Sugestoes ----------
add_heading_blue(doc, "3. Sugestões estratégicas adicionais", level=1)
add_bullets(doc, [
    "Não eliminar Gnomon, recontextualizá-lo: pode-se citar o Gnomon como motivação "
    "inicial e o contexto que levou a expandir o foco para ODS municipais via Google "
    "Places — preserva o histórico do projeto, mas o objeto principal fica claro.",
    "Ampliar o TAM: 4 respondentes é arriscado para a banca questionar quantos foram "
    "os avaliadores. Se conseguir mais 4–6 respostas em maio/junho, fortalece muito.",
    "Tabela comparativa de ODS no texto: criar uma tabela única consolidando ODS "
    "analisados, cidades aplicadas, nº de análises e % de problemas — vira referência "
    "visual forte para a banca.",
    "Gravar uma demo (vídeo curto) do formulario_config_ods.html gerando um config novo "
    "+ rodando ods_analise_generica.py — pode ser citado nos apêndices como link e usado "
    "na defesa.",
    "Submeter um segundo paper (revista) durante a janela de revisão das semanas 9–10 — "
    "o conteúdo do TAM + generalização para múltiplos ODS é novidade em relação ao EATIS "
    "e aproveita o material já escrito.",
])

doc.add_paragraph()

# ---------- 4. Marcos ----------
add_heading_blue(doc, "4. Marcos principais (resumo)", level=1)
mk_headers = ["Marco", "Data alvo"]
mk_rows = [
    ["Capítulos 1–4 (cabeçalho do trabalho) realinhados ao novo escopo", "07/05/2026"],
    ["Fundamentação Teórica + Trabalhos Relacionados concluídos", "14/05/2026"],
    ["Capítulo 7 (Implementação) atualizado com ODS 3/4 e generalização", "21/05/2026"],
    ["Capítulo 8 (Resultados) finalizado para todos os ODS e cidades", "04/06/2026"],
    ["Capítulo de Validação TAM finalizado", "11/06/2026"],
    ["Capítulo 9 (Discussão) finalizado", "18/06/2026"],
    ["Capítulo 10 (Conclusão) + referências + listas finalizados", "25/06/2026"],
    ["Revisão integral concluída", "02/07/2026"],
    ["Versão final entregue ao orientador", "06/07/2026"],
    ["Depósito da dissertação", "~15/07/2026"],
    ["Defesa", "~25/07–01/08/2026"],
]
make_table(doc, mk_headers, mk_rows, widths_cm=[12.0, 4.0])

doc.add_paragraph()

# ---------- assinatura ----------
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = foot.add_run("Mossoró, 30 de abril de 2026.")
run.italic = True
run.font.name = "Calibri"
run.font.size = Pt(10)

# ---------- salvar ----------
out_path = r"C:\Users\João Neto\Documents\Workspace\ppgcctex\Cronograma_Defesa_Dissertacao.docx"
doc.save(out_path)
print(f"Arquivo gerado em: {out_path}")
